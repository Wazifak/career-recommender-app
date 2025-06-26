# Innovigyan AI Career Recommender - App.py (Clean Version with Email and Analytics Setup)

import streamlit as st
import pandas as pd
import pickle
from docx import Document
from xhtml2pdf import pisa
import io

# Set page title and icon
st.set_page_config(page_title="Innovigyan AI Career Recommender", page_icon="🎓")

# Utility: Convert HTML to PDF
def create_pdf_from_html(html_content):
    result = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html_content), dest=result)
    return result.getvalue()

# Utility: Generate Word Doc
def create_word_doc(name, career, confidence, skills, interests):
    doc = Document()
    doc.add_heading('Career Recommendation Report', 0)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Recommended Career: {career}")
    doc.add_paragraph(f"Confidence: {confidence:.2f}%")
    doc.add_paragraph(f"Skills: {', '.join(skills)}")
    doc.add_paragraph(f"Interests: {', '.join(interests)}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# Load model and encoders
model = pickle.load(open("career_recommendation_model.pkl", "rb"))
encoders = pickle.load(open("career_encoders.pkl", "rb"))

# Load training dataset for visuals
df = pd.read_csv("AI-based Career Recommendation System.csv")

# App title and user input
st.title("🎓 Innovigyan AI Career Recommender")
name = st.text_input("👤 Enter your full name")

if name:
    clean_name = name.strip().title()
    st.markdown(f"## 👋 Welcome, **{clean_name}**!")
    st.markdown("Let AI guide your future based on your skills and interests.")
else:
    st.warning("Please enter your name to continue.")

# Sidebar inputs
st.sidebar.header("📋 Your Profile Details")
age = st.sidebar.slider("Select your age", 16, 60, 22)
education = st.sidebar.selectbox("Your highest education level", encoders["Education"].classes_)
skills_input = st.sidebar.multiselect("Select your skills", encoders["Skills"].classes_)
interests_input = st.sidebar.multiselect("Select your interests", encoders["Interests"].classes_)

# Predict button
if st.button("🔍 Recommend My Career"):
    if not name or not skills_input or not interests_input:
        st.warning("Please fill in all details to get your recommendation.")
    else:
        try:
            edu_encoded = encoders["Education"].transform([education])[0]
            skills_vector = [1 if skill in skills_input else 0 for skill in encoders["Skills"].classes_]
            interests_vector = [1 if interest in interests_input else 0 for interest in encoders["Interests"].classes_]

            input_data = pd.DataFrame(
                [[age, edu_encoded] + skills_vector + interests_vector],
                columns=["Age", "Education"] + list(encoders["Skills"].classes_) + list(encoders["Interests"].classes_)
            )

            proba = model.predict_proba(input_data)[0]
            top_3_idx = proba.argsort()[-3:][::-1]
            top_3_careers = encoders["Recommended_Career"].inverse_transform(top_3_idx)
            top_3_scores = proba[top_3_idx] * 100

            result = top_3_careers[0]
            confidence = top_3_scores[0]

            st.success(f"🎯 Recommended Career Path: **{result}**")
            st.info(f"📈 Match Confidence: {confidence:.2f}%")

            # Report download
            st.markdown("### 📥 Download Your Career Report")
            html_report = f"""
            <h2>Career Recommendation Report</h2>
            <p><b>Name:</b> {clean_name}</p>
            <p><b>Recommended Career:</b> {result}</p>
            <p><b>Confidence:</b> {confidence:.2f}%</p>
            <p><b>Skills:</b> {', '.join(skills_input)}</p>
            <p><b>Interests:</b> {', '.join(interests_input)}</p>
            """
            pdf_bytes = create_pdf_from_html(html_report)
            st.download_button("📄 Download PDF", pdf_bytes, file_name="career_report.pdf", mime='application/pdf')

            word_bytes = create_word_doc(clean_name, result, confidence, skills_input, interests_input)
            st.download_button("📝 Download Word", word_bytes, file_name="career_report.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            # Summary and visual insights
            st.subheader("📌 Top 3 Career Matches")
            chart_df = pd.DataFrame({'Career': top_3_careers, 'Confidence (%)': top_3_scores})
            st.bar_chart(chart_df.set_index("Career"))

            st.subheader("📋 Career Match Explanation")
            explanation = pd.DataFrame({
                "Career": top_3_careers,
                "Confidence": [f"{c:.2f}%" for c in top_3_scores],
                "Reason": [
                    "Strong match with your skill & interest set",
                    "Great fit with educational background",
                    "Alternative path based on overlapping traits"
                ]
            })
            st.dataframe(explanation)

        except Exception as e:
            st.error(f"❌ Error: {e}")

# Footer
st.markdown("---")
st.markdown("🔗 [Connect with me on LinkedIn](https://www.linkedin.com/in/wazifa-kapdi/)")
st.caption("© 2025 Wazifa Kapdi | All rights reserved.")
